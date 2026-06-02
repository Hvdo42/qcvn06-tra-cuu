"""
merge_content.py – Tạo QCVN06_Content.md từ:
  1. QCVN_06_2022_BXD.docx  (nội dung gốc đầy đủ, convert từ .doc)
  2. QCVN_06_2022_BXD_Hop_Nhat.docx (bản hợp nhất, có [SĐ1] markers)

Chiến lược:
  - Dùng file .docx gốc làm nguồn nội dung (đầy đủ, chính xác)
  - Dùng file hợp nhất để lấy danh sách clause có [SĐ1]
  - Khi ghi markdown: nếu clause nằm trong danh sách [SĐ1] → thêm marker
"""
import re, sys, json
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

# ===== ĐƯỜNG DẪN =====
MAIN_DOCX  = r'd:\Cowork\khoi luong\QC06\QCVN_06_2022_BXD.docx'
MERGED_DOCX= r'd:\Cowork\khoi luong\QC06\QCVN_06_2022_BXD_Hop_Nhat.docx'
MD_OUT     = r'd:\Cowork\khoi luong\QC06\QCVN06_Content.md'

# ===== REGEX =====
CLAUSE_RE = re.compile(
    r'^(\d+\.\d+(?:\.\d+)*[a-z]?'
    r'|[A-IĐ]\.\d+(?:\.\d+)*'
    r'|B[aả]ng\s+\d+'
    r'|\d+\.\d+[÷\-]\d+\.\d+)',
    re.IGNORECASE
)
SD1_RE = re.compile(r'\[S[DĐ]1', re.IGNORECASE)

# ===== BƯỚC 1: Lấy danh sách clause có [SĐ1] từ file hợp nhất =====
print('Đọc file hợp nhất để lấy [SĐ1] markers...')
sd1_clauses = set()
merged_doc  = Document(MERGED_DOCX)
for p in merged_doc.paragraphs:
    text = p.text.strip()
    if not text or not SD1_RE.search(text):
        continue
    m = CLAUSE_RE.match(text)
    if m:
        clause_id = m.group(1).strip().lower()
        sd1_clauses.add(clause_id)

# Thêm các clause từ nội dung bảng trong hợp nhất
for tbl in merged_doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if SD1_RE.search(text):
                m = CLAUSE_RE.match(text)
                if m:
                    sd1_clauses.add(m.group(1).strip().lower())

print(f'  → {len(sd1_clauses)} clauses có [SĐ1]: {sorted(sd1_clauses)[:10]}...')

# ===== BƯỚC 2: Đọc và parse file gốc đầy đủ =====
print('Đọc file gốc đầy đủ...')
main_doc = Document(MAIN_DOCX)

# Map style → cấp độ
STYLE_LEVEL = {
    'Heading 1': 1,   # Dùng cho PHẦN 6, 7 trong file này
    'Heading 2': 2,
    'Heading 3': 3,
    'Heading 4': 4,
    'Heading 5': 5,
    'Heading 6': 6,
    'PL H1': 3,    # Appendix sub-heading
    'PL H2': 3,
    'PL H3': 4,
    'PL H4': 4,
}

# Regex nhận diện PHỤ LỤC A-I ở đầu đoạn (dùng kết hợp với text pattern)
APPENDIX_HEADING_RE = re.compile(
    r'^PH[Ụ\w]\s*L[Ụ\w]C\s+([A-I])\b',
    re.IGNORECASE
)

def detect_appendix_from_text(text):
    """Nhận diện PHỤ LỤC A-I từ text (bất kể style)."""
    m = APPENDIX_HEADING_RE.match(text.strip())
    if m:
        k = m.group(1).upper()
        appendix_labels = {
            'A': 'PHỤ LỤC A – QUY ĐỊNH BỔ SUNG ĐỐI VỚI MỘT SỐ NHÓM NHÀ CỤ THỂ',
            'B': 'PHỤ LỤC B – PHÂN LOẠI VẬT LIỆU XÂY DỰNG THEO ĐẶC TÍNH KỸ THUẬT VỀ CHÁY',
            'C': 'PHỤ LỤC C – HẠNG NGUY HIỂM CHÁY VÀ CHÁY NỔ',
            'D': 'PHỤ LỤC D – BẢO VỆ CHỐNG KHÓI',
            'E': 'PHỤ LỤC E – KHOẢNG CÁCH PHÒNG CHÁY CHỐNG CHÁY',
            'F': 'PHỤ LỤC F – GIỚI HẠN CHỊU LỬA DANH ĐỊNH MỘT SỐ CẤU KIỆN',
            'G': 'PHỤ LỤC G – KHOẢNG CÁCH ĐẾN LỐI RA THOÁT NẠN VÀ CHIỀU RỘNG LỐI RA THOÁT NẠN',
            'H': 'PHỤ LỤC H – BẬC CHỊU LỬA VÀ CÁC YÊU CẦU BẢO ĐẢM AN TOÀN CHÁY',
            'I': 'PHỤ LỤC I – CÁC HÌNH MINH HỌA (Tham khảo)',
        }
        return appendix_labels.get(k)
    return None

# Section map từ heading text
def detect_main_section(text):
    """Nhận diện phần chính từ heading text."""
    t = text.strip().upper()
    # Phần chính: "1 QUY ĐỊNH CHUNG", "2 PHÂN LOẠI..."
    m = re.match(r'^([1-7])\s+(.+)', t)
    if m:
        num = int(m.group(1))
        labels = {
            1: 'PHẦN 1 – QUY ĐỊNH CHUNG',
            2: 'PHẦN 2 – PHÂN LOẠI KỸ THUẬT VỀ CHÁY',
            3: 'PHẦN 3 – BẢO ĐẢM AN TOÀN CHO NGƯỜI',
            4: 'PHẦN 4 – NGĂN CHẶN CHÁY LAN',
            5: 'PHẦN 5 – CẤP NƯỚC CHỮA CHÁY',
            6: 'PHẦN 6 – CHỮA CHÁY VÀ CỨU NẠN',
            7: 'PHẦN 7 – TỔ CHỨC THỰC HIỆN',
        }
        return labels.get(num)
    # Phụ lục: "PHỤ LỤC A..."
    m2 = re.match(r'^PH[Ụ\w]\s*L[Ụ\w]C\s+([A-I])\b', t)
    if m2:
        k = m2.group(1)
        appendix_labels = {
            'A': 'PHỤ LỤC A – QUY ĐỊNH BỔ SUNG ĐỐI VỚI MỘT SỐ NHÓM NHÀ CỤ THỂ',
            'B': 'PHỤ LỤC B – PHÂN LOẠI VẬT LIỆU XÂY DỰNG THEO ĐẶC TÍNH KỸ THUẬT VỀ CHÁY',
            'C': 'PHỤ LỤC C – HẠNG NGUY HIỂM CHÁY VÀ CHÁY NỔ',
            'D': 'PHỤ LỤC D – BẢO VỆ CHỐNG KHÓI',
            'E': 'PHỤ LỤC E – KHOẢNG CÁCH PHÒNG CHÁY CHỐNG CHÁY',
            'F': 'PHỤ LỤC F – GIỚI HẠN CHỊU LỬA DANH ĐỊNH MỘT SỐ CẤU KIỆN',
            'G': 'PHỤ LỤC G – KHOẢNG CÁCH ĐẾN LỐI RA THOÁT NẠN VÀ CHIỀU RỘNG LỐI RA THOÁT NẠN',
            'H': 'PHỤ LỤC H – BẬC CHỊU LỬA VÀ CÁC YÊU CẦU BẢO ĐẢM AN TOÀN CHÁY',
            'I': 'PHỤ LỤC I – CÁC HÌNH MINH HỌA (Tham khảo)',
        }
        return appendix_labels.get(k)
    return None

def get_sd1_status(text):
    """Kiểm tra clause trong text có [SĐ1] không (từ file hợp nhất)."""
    m = CLAUSE_RE.match(text.strip())
    if not m:
        return False
    clause_id = m.group(1).strip().lower()
    return clause_id in sd1_clauses

# ===== BƯỚC 3: Tạo markdown =====
print('Tạo markdown...')
lines_out = []
lines_out.append('# QCVN 06:2022/BXD – An toàn cháy cho nhà và công trình')
lines_out.append('')
lines_out.append('> **Nguồn:** Từ file Word gốc QCVN 06:2022/BXD (Thông tư 06/2022/TT-BXD).')
lines_out.append('> **Sửa đổi:** Các điều khoản đánh dấu [SĐ1] đã được sửa đổi/bổ sung theo Thông tư 09/2023/TT-BXD.')
lines_out.append('> Sau khi chỉnh sửa file này, chạy `python rebuild.py` để cập nhật app HTML.')
lines_out.append('')
lines_out.append('---')
lines_out.append('')

current_main_section = None
skip_until_section = True  # Bỏ qua phần đầu (thông tư, mục lục)
stats = {'total': 0, 'sd1_marked': 0, 'sections': 0}

for p in main_doc.paragraphs:
    text = p.text.strip()
    if not text:
        continue

    style_name = p.style.name if p.style else 'Normal'
    level = STYLE_LEVEL.get(style_name, 0)

    # Phát hiện section chính từ Heading 1 hoặc Heading 2
    if style_name in ('Heading 1', 'Heading 2') and level in (1, 2):
        main_sec = detect_main_section(text)
        if main_sec:
            if skip_until_section:
                skip_until_section = False
            if main_sec != current_main_section:
                current_main_section = main_sec
                lines_out.append('')
                lines_out.append(f'# {main_sec}')
                lines_out.append('')
                stats['sections'] += 1
                continue  # Đã ghi section label, không ghi heading riêng

    # Phát hiện PHỤ LỤC từ text (bất kể style - vì dùng nhiều style khác nhau)
    if not skip_until_section:
        appendix_label = detect_appendix_from_text(text)
        if appendix_label and appendix_label != current_main_section:
            current_main_section = appendix_label
            lines_out.append('')
            lines_out.append(f'# {appendix_label}')
            lines_out.append('')
            stats['sections'] += 1
            continue

    # Bỏ qua phần trước khi vào phần chính (thông tư, mục lục, lời nói đầu)
    if skip_until_section:
        continue

    # Xác định [SĐ1] status
    is_sd1 = get_sd1_status(text)
    sd1_marker = ' [SĐ1]' if is_sd1 else ''
    if is_sd1:
        stats['sd1_marked'] += 1

    # Format theo style
    if level in (2, 3):
        # Heading: in đậm với dấu ##
        lines_out.append(f'## {text}{sd1_marker}')
        lines_out.append('')
    elif level in (4, 5, 6):
        # Sub-item
        lines_out.append(f'  {text}{sd1_marker}')
        lines_out.append('')
    elif style_name in ('CHUTHICH', 'List Bullet', 'List Bullet 2', 'List Bullet 3'):
        # Chú thích hoặc bullet list
        lines_out.append(f'  > {text}')
        lines_out.append('')
    elif style_name == 'Caption':
        # Tên bảng/hình
        lines_out.append(f'**{text}**')
        lines_out.append('')
    elif style_name in ('Body Text', 'Body Text 2', 'Body Text 3', 'Normal', 'Table just'):
        # Nội dung thông thường — kiểm tra có phải clause heading không
        m = CLAUSE_RE.match(text)
        if m and style_name in ('Body Text', 'Normal'):
            # Có thể là clause với style Body Text (xảy ra đôi khi)
            clause_sd1 = get_sd1_status(text)
            sd1_m = ' [SĐ1]' if clause_sd1 else ''
            lines_out.append(f'{text}{sd1_m}')
        else:
            lines_out.append(text)
        lines_out.append('')
    else:
        # Style khác — ghi bình thường
        lines_out.append(text)
        lines_out.append('')

    stats['total'] += 1

# Xử lý bảng (tables)
for tbl in main_doc.tables:
    # Bỏ qua bảng nếu chưa vào phần chính — khó xác định vị trí chính xác
    # Ghi tóm tắt bảng (header row)
    if tbl.rows:
        header = ' | '.join(c.text.strip() for c in tbl.rows[0].cells if c.text.strip())
        if header:
            pass  # Tables được xử lý trong luồng paragraph bởi Word COM, bỏ qua ở đây

# Loại bỏ dòng trống liên tiếp
result_lines = []
prev_empty = False
for line in lines_out:
    is_empty = not line.strip()
    if is_empty and prev_empty:
        continue
    result_lines.append(line)
    prev_empty = is_empty

output = '\n'.join(result_lines)
with open(MD_OUT, 'w', encoding='utf-8') as f:
    f.write(output)

import os
size_kb = os.path.getsize(MD_OUT) / 1024
print(f'\nKết quả:')
print(f'  Tổng entries: {stats["total"]}')
print(f'  Sections: {stats["sections"]}')
print(f'  Clauses có [SĐ1]: {stats["sd1_marked"]}')
print(f'  File: {MD_OUT} ({size_kb:.0f} KB)')
print(f'\nXong! Mở QCVN06_Content.md để kiểm tra.')
